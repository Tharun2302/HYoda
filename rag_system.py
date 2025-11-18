"""
RAG System for HealthYoda Question Book
Extracts questions and answers from the .docx file and provides retrieval capabilities
"""
import docx
import os
from typing import List, Dict, Optional, Tuple
import json

class QuestionBookRAG:
    """RAG system for retrieving relevant questions from the Question Book"""
    
    def __init__(self, docx_path: str = 'docx/Question BOOK.docx'):
        self.docx_path = docx_path
        self.questions = []
        self.sections = []
        self.load_document()
    
    def load_document(self):
        """Load and parse the .docx document"""
        if not os.path.exists(self.docx_path):
            raise FileNotFoundError(f"Question book not found at {self.docx_path}")
        
        doc = docx.Document(self.docx_path)
        current_system = None
        current_symptom = None
        current_category = None
        current_question = None
        current_answers = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            if not text:
                continue
            
            # Detect system headers
            if any(system in text for system in ['Cardiac System', 'Respiratory System', 'GI System', 
                                                 'Neurologic System', 'Musculoskeletal System', 'GU System',
                                                 'Dermatologic System', 'Endocrine', 'ENT Eye System']):
                if 'System' in text and len(text) < 100:
                    current_system = text
                    current_symptom = None
                    current_category = None
            
            # Detect symptom/complaint headers (usually bold or standalone lines)
            elif text and not text.startswith('Q:') and not text.startswith('Possible') and \
                 not text.startswith('-') and len(text) < 100 and \
                 text not in ['Chief Complaint', 'Onset/Duration', 'Quality/Severity', 
                             'Aggravating/Relieving', 'Associated Symptoms', 'Red Flags', 
                             'ROS', 'Context', 'Table of Contents']:
                # Likely a symptom/complaint name
                if current_system and text != current_system:
                    current_symptom = text
                    current_category = None
            
            # Detect category headers
            elif text in ['Chief Complaint', 'Onset/Duration', 'Quality/Severity', 
                         'Aggravating/Relieving', 'Associated Symptoms', 'Red Flags', 
                         'ROS', 'Context']:
                current_category = text
                current_question = None
                current_answers = []
            
            # Detect questions
            elif text.startswith('Q:') or text.startswith('Q.'):
                # Save previous question if exists
                if current_question and current_system:
                    # Create tags for tracking
                    tags = []
                    if current_system:
                        tags.append(f"system:{current_system}")
                    if current_symptom:
                        tags.append(f"symptom:{current_symptom}")
                    if current_category:
                        tags.append(f"category:{current_category}")
                    
                    # Create question tree path
                    tree_path = " > ".join(filter(None, [
                        current_system,
                        current_symptom,
                        current_category
                    ]))
                    
                    self.questions.append({
                        'system': current_system,
                        'symptom': current_symptom,
                        'category': current_category,
                        'question': current_question,
                        'possible_answers': current_answers.copy(),
                        'line_number': i,
                        'tags': tags,
                        'tree_path': tree_path
                    })
                
                current_question = text.replace('Q:', '').replace('Q.', '').strip()
                current_answers = []
            
            # Detect possible answers
            elif text.startswith('Possible Answers:') or text.startswith('Possible answers:'):
                current_answers = []
            
            elif text.startswith('-') and current_question:
                answer = text[1:].strip()
                if answer:
                    current_answers.append(answer)
        
        # Save last question
        if current_question and current_system:
            tags = []
            if current_system:
                tags.append(f"system:{current_system}")
            if current_symptom:
                tags.append(f"symptom:{current_symptom}")
            if current_category:
                tags.append(f"category:{current_category}")
            
            tree_path = " > ".join(filter(None, [
                current_system,
                current_symptom,
                current_category
            ]))
            
            self.questions.append({
                'system': current_system,
                'symptom': current_symptom,
                'category': current_category,
                'question': current_question,
                'possible_answers': current_answers.copy(),
                'line_number': len(doc.paragraphs),
                'tags': tags,
                'tree_path': tree_path
            })
        
        print(f"Loaded {len(self.questions)} questions from {len(set(q['system'] for q in self.questions))} systems")
    
    def search_by_system(self, system_name: str) -> List[Dict]:
        """Retrieve all questions for a specific system"""
        return [q for q in self.questions if system_name.lower() in q['system'].lower()]
    
    def search_by_symptom(self, symptom: str) -> List[Dict]:
        """Retrieve questions for a specific symptom/complaint"""
        symptom_lower = symptom.lower()
        return [q for q in self.questions 
                if q['symptom'] and symptom_lower in q['symptom'].lower()]
    
    def search_by_category(self, category: str) -> List[Dict]:
        """Retrieve questions for a specific category (e.g., 'Onset/Duration', 'Red Flags')"""
        return [q for q in self.questions if q['category'] == category]
    
    def search_by_keywords(self, keywords: List[str], system: Optional[str] = None) -> List[Dict]:
        """Search questions by keywords in question text"""
        results = []
        keywords_lower = [k.lower() for k in keywords]
        
        for q in self.questions:
            if system and system.lower() not in q['system'].lower():
                continue
            
            question_lower = q['question'].lower()
            if any(kw in question_lower for kw in keywords_lower):
                results.append(q)
        
        return results
    
    def get_next_question(self, conversation_context: str, current_category: Optional[str] = None,
                         symptom: Optional[str] = None, system: Optional[str] = None) -> Optional[Dict]:
        """
        Get the next relevant question based on conversation context
        Returns question with tags and tree_path for logging
        """
        context_lower = conversation_context.lower()
        
        # Determine which system/symptom we're discussing
        systems_keywords = {
            'cardiac': ['chest', 'heart', 'cardiac', 'palpitation', 'shortness of breath'],
            'respiratory': ['breathing', 'cough', 'respiratory', 'lung', 'wheeze'],
            'gi': ['stomach', 'abdominal', 'nausea', 'vomiting', 'diarrhea', 'gi', 'gastro'],
            'neurologic': ['headache', 'dizziness', 'seizure', 'neurologic', 'neurological'],
            'musculoskeletal': ['joint', 'muscle', 'bone', 'pain', 'musculoskeletal'],
            'gu': ['urinary', 'bladder', 'kidney', 'gu', 'genitourinary'],
            'dermatologic': ['skin', 'rash', 'dermatologic', 'dermatological'],
            'endocrine': ['diabetes', 'thyroid', 'endocrine', 'hormone'],
            'ent': ['ear', 'nose', 'throat', 'ent', 'hearing', 'vision']
        }
        
        detected_system = system
        if not detected_system:
            for sys_name, keywords in systems_keywords.items():
                if any(kw in context_lower for kw in keywords):
                    detected_system = sys_name
                    break
        
        # Filter questions
        candidates = self.questions
        
        if detected_system:
            candidates = [q for q in candidates if detected_system.lower() in q['system'].lower()]
        
        if symptom:
            candidates = [q for q in candidates if q['symptom'] and symptom.lower() in q['symptom'].lower()]
        
        if current_category:
            candidates = [q for q in candidates if q['category'] == current_category]
        
        # Return first relevant question with tags
        if candidates:
            question = candidates[0].copy()
            # Ensure tags and tree_path exist
            if 'tags' not in question:
                tags = []
                if question.get('system'):
                    tags.append(f"system:{question['system']}")
                if question.get('symptom'):
                    tags.append(f"symptom:{question['symptom']}")
                if question.get('category'):
                    tags.append(f"category:{question['category']}")
                question['tags'] = tags
                question['tree_path'] = " > ".join(filter(None, [
                    question.get('system'),
                    question.get('symptom'),
                    question.get('category')
                ]))
            return question
        
        return None
    
    def get_all_categories_for_symptom(self, symptom: str) -> List[str]:
        """Get all question categories available for a symptom"""
        questions = self.search_by_symptom(symptom)
        categories = list(set(q['category'] for q in questions if q['category']))
        return categories
    
    def get_questions_by_phase(self, phase: str) -> List[Dict]:
        """
        Get questions based on intake phase
        Maps intake phases to question categories
        """
        phase_mapping = {
            'greeting': ['Chief Complaint'],
            'symptom_discovery': ['Chief Complaint', 'Onset/Duration', 'Quality/Severity', 
                                 'Aggravating/Relieving', 'Associated Symptoms'],
            'red_flags': ['Red Flags'],
            'review_of_systems': ['ROS'],
            'context': ['Context']
        }
        
        categories = phase_mapping.get(phase.lower(), [])
        return [q for q in self.questions if q['category'] in categories]
    
    def export_to_json(self, output_path: str = 'question_book_data.json'):
        """Export parsed questions to JSON for easier access"""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.questions, f, indent=2, ensure_ascii=False)
        print(f"Exported {len(self.questions)} questions to {output_path}")

# Test the RAG system
if __name__ == '__main__':
    rag = QuestionBookRAG()
    
    print("\n" + "="*80)
    print("RAG SYSTEM TEST")
    print("="*80)
    
    # Test searches
    print("\n1. Cardiac system questions:")
    cardiac_qs = rag.search_by_system('Cardiac')
    print(f"   Found {len(cardiac_qs)} questions")
    if cardiac_qs:
        print(f"   Sample: {cardiac_qs[0]['question']}")
    
    print("\n2. Chest pain questions:")
    chest_pain_qs = rag.search_by_symptom('Chest Pain')
    print(f"   Found {len(chest_pain_qs)} questions")
    
    print("\n3. Red flag questions:")
    red_flag_qs = rag.search_by_category('Red Flags')
    print(f"   Found {len(red_flag_qs)} questions")
    if red_flag_qs:
        print(f"   Sample: {red_flag_qs[0]['question']}")
        print(f"   Possible answers: {red_flag_qs[0]['possible_answers'][:3]}")
    
    print("\n4. Questions for symptom discovery phase:")
    symptom_qs = rag.get_questions_by_phase('symptom_discovery')
    print(f"   Found {len(symptom_qs)} questions")
    
    # Export to JSON
    rag.export_to_json()

